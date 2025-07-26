#!/usr/bin/env python3
"""
Unified Data Analysis Tool
=========================

A comprehensive tool that combines file extraction, LLM analysis, structure templating,
and professional report generation into a single unified application.

Features:
- Interactive file selection with preview
- Multiple file format support (CSV, JSON, TXT, MD, DOCX, etc.)
- Structure template extraction for consistent formatting
- LLM integration (OpenAI, Anthropic, Ollama)
- Professional DOCX report generation
- Enhanced table detection and formatting
- User-friendly command-line interface

Usage:
    python data_analysis_tool.py

Author: Data Analysis Team
Version: 2.0
"""

import os
import sys
import json
import csv
import re
import requests
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple

# Try to import optional dependencies
try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class FileExtractor:
    """Handles file content extraction with enhanced table detection."""
    
    def __init__(self):
        self.supported_extensions = [
            '.txt', '.log', '.md', '.py', '.js', '.html', '.css', '.xml',
            '.json', '.csv', '.yml', '.yaml', '.ini', '.conf', '.tsv'
        ]
    
    def extract_file_content(self, file_path):
        """Extract content from various file types with enhanced table detection."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return f"Error: File '{file_path}' does not exist."
            
            extension = file_path.suffix.lower()
            
            # Text files
            if extension in ['.txt', '.md', '.csv', '.tsv', '.log']:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    content = file.read()
                    
                    if extension in ['.csv', '.tsv'] or self.detect_tabular_data(content):
                        return self.format_tabular_content(content, extension)
                    
                    return content
            
            # JSON files
            elif extension == '.json':
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                    formatted_content = json.dumps(data, indent=2, ensure_ascii=False)
                    
                    if self.is_json_tabular(data):
                        formatted_content += "\n\n" + self.format_json_as_table(data)
                    
                    return formatted_content
            
            # XML files
            elif extension == '.xml':
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                    return file.read()
            
            # PDF files
            elif extension == '.pdf' and PDF_AVAILABLE:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        content = ""
                        for page in pdf_reader.pages:
                            content += page.extract_text() + "\n"
                        return content
                except Exception:
                    return "Error: Could not extract PDF content."
            
            # Word documents
            elif extension in ['.docx', '.doc'] and DOCX_AVAILABLE:
                try:
                    doc = Document(file_path)
                    content = ""
                    
                    for paragraph in doc.paragraphs:
                        content += paragraph.text + "\n"
                    
                    if doc.tables:
                        content += "\n=== TABLES FOUND ===\n"
                        for i, table in enumerate(doc.tables):
                            content += f"\nTable {i+1}:\n"
                            content += self.format_docx_table(table)
                    
                    return content
                except Exception:
                    return "Error: Could not extract Word document content."
            
            # Excel files
            elif extension in ['.xlsx', '.xls'] and PANDAS_AVAILABLE:
                try:
                    excel_file = pd.ExcelFile(file_path)
                    content = ""
                    
                    for sheet_name in excel_file.sheet_names:
                        df = pd.read_excel(file_path, sheet_name=sheet_name)
                        content += f"\n=== SHEET: {sheet_name} ===\n"
                        content += self.format_dataframe_as_table(df)
                        content += "\n"
                    
                    return content
                except Exception:
                    return "Error: Could not extract Excel content."
            
            # Default: try to read as text
            else:
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                        content = file.read()
                        
                        if self.detect_tabular_data(content):
                            return self.format_tabular_content(content, extension)
                        
                        return content
                except Exception as e:
                    return f"Error reading file: {str(e)}"
                    
        except Exception as e:
            return f"Error extracting content from '{file_path}': {str(e)}"
    
    def detect_tabular_data(self, content):
        """Detect if content contains tabular data."""
        lines = content.strip().split('\n')
        if len(lines) < 2:
            return False
        
        delimiters = [',', '\t', '|', ';']
        
        for delimiter in delimiters:
            first_line_count = lines[0].count(delimiter)
            if first_line_count > 0:
                consistent_count = 0
                for line in lines[1:min(5, len(lines))]:
                    if abs(line.count(delimiter) - first_line_count) <= 1:
                        consistent_count += 1
                
                if consistent_count >= min(2, len(lines) - 1):
                    return True
        
        return False
    
    def format_tabular_content(self, content, file_extension):
        """Format tabular content for better analysis."""
        lines = content.strip().split('\n')
        
        delimiter = ','
        if file_extension == '.tsv':
            delimiter = '\t'
        elif '|' in lines[0]:
            delimiter = '|'
        elif ';' in lines[0] and lines[0].count(';') > lines[0].count(','):
            delimiter = ';'
        
        formatted_content = f"TABULAR DATA DETECTED (Delimiter: '{delimiter}')\n"
        formatted_content += "=" * 50 + "\n\n"
        formatted_content += content
        formatted_content += "\n\n" + "=" * 50
        formatted_content += "\nDATA ANALYSIS NOTES:"
        formatted_content += f"\n- Total rows: {len(lines)}"
        formatted_content += f"\n- Delimiter used: '{delimiter}'"
        
        if len(lines) > 0:
            columns = len(lines[0].split(delimiter))
            formatted_content += f"\n- Estimated columns: {columns}"
            
            if len(lines) > 1:
                formatted_content += f"\n- Header row: {lines[0]}"
                formatted_content += f"\n- Sample data row: {lines[1] if len(lines) > 1 else 'N/A'}"
        
        return formatted_content
    
    def format_docx_table(self, table):
        """Format a Word table as text."""
        content = ""
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            content += " | ".join(row_data) + "\n"
        return content
    
    def format_dataframe_as_table(self, df):
        """Format pandas DataFrame as markdown table."""
        if df.empty:
            return "No data in this sheet."
        
        content = "| " + " | ".join(str(col) for col in df.columns) + " |\n"
        content += "|" + "|".join(["-" * 10] * len(df.columns)) + "|\n"
        
        for idx, row in df.head(10).iterrows():
            content += "| " + " | ".join(str(val) for val in row.values) + " |\n"
        
        if len(df) > 10:
            content += f"\n... and {len(df) - 10} more rows\n"
        
        content += f"\nDataFrame Info:\n"
        content += f"- Shape: {df.shape[0]} rows × {df.shape[1]} columns\n"
        content += f"- Columns: {', '.join(df.columns.astype(str))}\n"
        
        return content
    
    def is_json_tabular(self, data):
        """Check if JSON data represents tabular information."""
        if isinstance(data, list) and len(data) > 0:
            if isinstance(data[0], dict):
                first_keys = set(data[0].keys())
                return all(isinstance(item, dict) and set(item.keys()) == first_keys for item in data[:5])
        return False
    
    def format_json_as_table(self, data):
        """Format JSON data as a table if it's tabular."""
        if not self.is_json_tabular(data):
            return ""
        
        headers = list(data[0].keys())
        
        content = "\nTABULAR VIEW:\n"
        content += "| " + " | ".join(headers) + " |\n"
        content += "|" + "|".join(["-" * 10] * len(headers)) + "|\n"
        
        for item in data[:10]:
            row_values = [str(item.get(header, '')) for header in headers]
            content += "| " + " | ".join(row_values) + " |\n"
        
        if len(data) > 10:
            content += f"\n... and {len(data) - 10} more records\n"
        
        return content
    
    def get_file_info(self, file_path):
        """Extract file metadata and content."""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": f"File '{file_path}' does not exist."}
        
        stats = file_path.stat()
        
        info = {
            "filename": file_path.name,
            "full_path": str(file_path.absolute()),
            "size_bytes": stats.st_size,
            "size_readable": self.format_file_size(stats.st_size),
            "modified_time": stats.st_mtime,
            "extension": file_path.suffix,
            "content": self.extract_file_content(file_path)
        }
        
        return info
    
    def format_file_size(self, size_bytes):
        """Convert bytes to human readable format."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024**2:
            return f"{size_bytes/1024:.1f} KB"
        elif size_bytes < 1024**3:
            return f"{size_bytes/1024**2:.1f} MB"
        else:
            return f"{size_bytes/1024**3:.1f} GB"


class StructurePreview:
    """Handles structure extraction and template management for consistent report formatting."""
    
    def __init__(self):
        self.structure_template = None
        self.preview_file_info = None
        self.extractor = FileExtractor()
    
    def extract_structure_from_file(self, file_path: str) -> Dict:
        """Extract structure template from a preview file."""
        try:
            content = self.extractor.extract_file_content(file_path)
            
            if not content:
                raise ValueError("Could not extract content from preview file")
            
            structure = self._analyze_content_structure(content)
            
            self.structure_template = structure
            self.preview_file_info = {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'file_size': os.path.getsize(file_path)
            }
            
            return structure
            
        except Exception as e:
            raise Exception(f"Error extracting structure from {file_path}: {str(e)}")
    
    def _analyze_content_structure(self, content: str) -> Dict:
        """Analyze content to extract structural elements."""
        structure = {
            'sections': [],
            'formatting_patterns': {},
            'content_types': [],
            'organization_style': 'default'
        }
        
        lines = content.split('\n')
        current_section = None
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
            
            section_info = self._detect_section(line, i)
            if section_info:
                if current_section:
                    structure['sections'].append(current_section)
                current_section = section_info
                continue
            
            content_type = self._detect_content_type(line)
            if content_type and current_section:
                if 'content_types' not in current_section:
                    current_section['content_types'] = []
                if content_type not in current_section['content_types']:
                    current_section['content_types'].append(content_type)
        
        if current_section:
            structure['sections'].append(current_section)
        
        structure['formatting_patterns'] = self._detect_formatting_patterns(content)
        structure['organization_style'] = self._determine_organization_style(structure['sections'])
        
        return structure
    
    def _detect_section(self, line: str, line_number: int) -> Optional[Dict]:
        """Detect if a line represents a section header."""
        # Numbered sections
        numbered_match = re.match(r'^(\d+)\.\s*(.+)', line)
        if numbered_match:
            return {
                'type': 'numbered_section',
                'number': numbered_match.group(1),
                'title': numbered_match.group(2),
                'line_number': line_number,
                'content_types': []
            }
        
        # Markdown headers
        header_match = re.match(r'^(#{1,6})\s*(.+)', line)
        if header_match:
            level = len(header_match.group(1))
            return {
                'type': 'markdown_header',
                'level': level,
                'title': header_match.group(2),
                'line_number': line_number,
                'content_types': []
            }
        
        # Uppercase headers
        if line.isupper() and len(line) > 3:
            return {
                'type': 'uppercase_header',
                'title': line,
                'line_number': line_number,
                'content_types': []
            }
        
        # Title case headers
        if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*:?\s*$', line) and len(line.split()) <= 5:
            return {
                'type': 'title_case_header',
                'title': line.rstrip(':'),
                'line_number': line_number,
                'content_types': []
            }
        
        return None
    
    def _detect_content_type(self, line: str) -> Optional[str]:
        """Detect the type of content in a line."""
        if '|' in line and line.count('|') >= 2:
            return 'table'
        
        if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
            return 'list'
        
        if line.startswith('```') or line.startswith('    '):
            return 'code'
        
        if line.startswith('>'):
            return 'quote'
        
        if ':' in line and len(line.split(':')) == 2:
            return 'key_value'
        
        if re.search(r'\d+%|\d+\.\d+|\$\d+|\d+\s*(units?|items?|records?)', line):
            return 'statistics'
        
        return 'paragraph'
    
    def _detect_formatting_patterns(self, content: str) -> Dict:
        """Detect formatting patterns in the content."""
        patterns = {
            'uses_markdown': False,
            'uses_numbering': False,
            'uses_bullets': False,
            'uses_tables': False,
            'uses_emphasis': False,
            'line_separators': False
        }
        
        if re.search(r'#{1,6}\s+', content) or re.search(r'\*\*.*?\*\*', content):
            patterns['uses_markdown'] = True
        
        if re.search(r'^\d+\.\s+', content, re.MULTILINE):
            patterns['uses_numbering'] = True
        
        if re.search(r'^\s*[-*+]\s+', content, re.MULTILINE):
            patterns['uses_bullets'] = True
        
        if '|' in content and content.count('|') >= 4:
            patterns['uses_tables'] = True
        
        if re.search(r'\*.*?\*|_.*?_|\*\*.*?\*\*', content):
            patterns['uses_emphasis'] = True
        
        if re.search(r'^[=\-_]{3,}$', content, re.MULTILINE):
            patterns['line_separators'] = True
        
        return patterns
    
    def _determine_organization_style(self, sections: List[Dict]) -> str:
        """Determine the overall organization style."""
        if not sections:
            return 'simple'
        
        numbered_sections = sum(1 for s in sections if s.get('type') == 'numbered_section')
        if numbered_sections == len(sections):
            return 'numbered'
        
        markdown_sections = sum(1 for s in sections if s.get('type') == 'markdown_header')
        if markdown_sections > len(sections) * 0.7:
            return 'markdown'
        
        if any(s.get('level', 0) > 1 for s in sections):
            return 'hierarchical'
        
        return 'mixed'
    
    def generate_structure_prompt(self) -> str:
        """Generate a prompt for the LLM based on the extracted structure."""
        if not self.structure_template:
            return ""
        
        prompt = f"""
STRUCTURE TEMPLATE (Based on preview file: {self.preview_file_info['file_name']}):

Please format your analysis following this exact structure:

"""
        
        if self.structure_template['sections']:
            prompt += "REQUIRED SECTIONS:\n"
            for i, section in enumerate(self.structure_template['sections'], 1):
                section_title = section['title']
                section_type = section['type']
                
                if section_type == 'numbered_section':
                    prompt += f"{i}. {section_title}\n"
                elif section_type == 'markdown_header':
                    level = section.get('level', 1)
                    prompt += f"{'#' * level} {section_title}\n"
                else:
                    prompt += f"{section_title}\n"
                
                if section.get('content_types'):
                    content_types = ', '.join(section['content_types'])
                    prompt += f"   (Include: {content_types})\n"
                prompt += "\n"
        
        formatting = self.structure_template['formatting_patterns']
        if any(formatting.values()):
            prompt += "FORMATTING REQUIREMENTS:\n"
            if formatting.get('uses_tables'):
                prompt += "- Include tables in markdown format when presenting data\n"
            if formatting.get('uses_bullets'):
                prompt += "- Use bullet points for lists\n"
            if formatting.get('uses_numbering'):
                prompt += "- Use numbered lists for sequential items\n"
            if formatting.get('uses_emphasis'):
                prompt += "- Use **bold** for emphasis and key points\n"
            prompt += "\n"
        
        prompt += """
IMPORTANT: 
- Follow the exact section structure shown above
- Maintain consistent formatting throughout
- Include all required content types in each section
- Ensure professional presentation matching the template style
"""
        
        return prompt
    
    def get_structure_summary(self) -> str:
        """Get a summary of the current structure template."""
        if not self.structure_template:
            return "No structure template loaded."
        
        summary = f"Structure Template from: {self.preview_file_info['file_name']}\n"
        summary += f"Organization Style: {self.structure_template['organization_style']}\n"
        summary += f"Sections: {len(self.structure_template['sections'])}\n"
        
        if self.structure_template['sections']:
            summary += "\nSection Overview:\n"
            for section in self.structure_template['sections']:
                summary += f"  - {section['title']} ({section['type']})\n"
        
        return summary


class LLMAnalyzer:
    """Handles LLM integration for file analysis."""
    
    def __init__(self, api_key=None, api_provider="openai"):
        self.api_key = api_key or os.getenv('API_KEY')
        self.api_provider = api_provider.lower()
        
        self.endpoints = {
            "openai": "https://api.openai.com/v1/chat/completions",
            "anthropic": "https://api.anthropic.com/v1/messages",
            "ollama": "http://localhost:11434/api/generate"
        }
    
    def analyze_content(self, content, file_path, structure_template=""):
        """Send content to LLM with analysis prompt and optional structure template."""
        file_name = Path(file_path).name
        
        system_prompt = f"""You are a data analyst. Analyze the following file content and create a comprehensive data report.

Instructions:
- Summarize the key information and insights from the data
- Identify patterns, trends, or notable findings
- Provide statistical analysis if applicable (counts, percentages, distributions)
- Highlight any data quality issues or anomalies
- Structure your response as a detailed professional report
- Include executive summary, detailed findings, and recommendations if relevant
- If the data contains tabular information, create summary tables using markdown format (| Column | Column |)
- For CSV data or structured data, provide key statistics in table format
- Include data distribution tables, top/bottom performers, or category breakdowns as appropriate
- Use tables to present numerical findings, comparisons, and statistical summaries clearly

Table Formatting Guidelines:
- Use markdown table syntax: | Header 1 | Header 2 | Header 3 |
- Separate headers with: |----------|----------|----------|
- Align numerical data properly
- Include totals, averages, or percentages where relevant
- Create multiple tables for different aspects of the analysis

File being analyzed: {file_name}
"""

        if structure_template:
            system_prompt += f"\n\n{structure_template}"

        user_prompt = f"Please analyze this file content and provide a detailed data report:\n\n{content}"
        
        if self.api_provider == "openai":
            return self._send_to_openai(system_prompt, user_prompt)
        elif self.api_provider == "anthropic":
            return self._send_to_anthropic(system_prompt, user_prompt)
        elif self.api_provider == "ollama":
            return self._send_to_ollama(system_prompt, user_prompt)
        else:
            raise ValueError(f"Unsupported API provider: {self.api_provider}")
    
    def _send_to_openai(self, system_prompt, user_prompt):
        """Send request to OpenAI API."""
        headers = {
            "Authorization": f"Bearer {self.api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "model": "gpt-4",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            "max_tokens": 2000,
            "temperature": 0.3
        }
        
        response = requests.post(self.endpoints["openai"], headers=headers, json=data)
        
        if response.status_code == 200:
            return response.json()["choices"][0]["message"]["content"]
        else:
            raise Exception(f"OpenAI API error: {response.status_code} - {response.text}")
    
    def _send_to_anthropic(self, system_prompt, user_prompt):
        """Send request to Anthropic Claude API."""
        headers = {
            "x-api-key": self.api_key,
            "content-type": "application/json",
            "anthropic-version": "2023-06-01"
        }
        
        data = {
            "model": "claude-3-sonnet-20240229",
            "max_tokens": 2000,
            "system": system_prompt,
            "messages": [
                {"role": "user", "content": user_prompt}
            ]
        }
        
        response = requests.post(self.endpoints["anthropic"], headers=headers, json=data)
        
        if response.status_code == 200:
            return response.json()["content"][0]["text"]
        else:
            raise Exception(f"Anthropic API error: {response.status_code} - {response.text}")
    
    def _send_to_ollama(self, system_prompt, user_prompt):
        """Send request to local Ollama instance."""
        combined_prompt = f"{system_prompt}\n\nUser Request: {user_prompt}"
        
        data = {
            "model": "llama2",
            "prompt": combined_prompt,
            "stream": False
        }
        
        response = requests.post(self.endpoints["ollama"], json=data)
        
        if response.status_code == 200:
            return response.json()["response"]
        else:
            raise Exception(f"Ollama API error: {response.status_code} - {response.text}")


class DocxGenerator:
    """Handles DOCX report generation with professional formatting."""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX generation. Install with: pip install python-docx")
        
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Set up custom styles for the document."""
        # Title style
        title_style = self.doc.styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_font = title_style.font
        title_font.name = 'Arial'
        title_font.size = Pt(18)
        title_font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
        
        # Heading style
        heading_style = self.doc.styles.add_style('CustomHeading', WD_STYLE_TYPE.PARAGRAPH)
        heading_font = heading_style.font
        heading_font.name = 'Arial'
        heading_font.size = Pt(14)
        heading_font.bold = True
        heading_style.paragraph_format.space_before = Pt(12)
        heading_style.paragraph_format.space_after = Pt(6)
        
        # Subheading style
        subheading_style = self.doc.styles.add_style('CustomSubheading', WD_STYLE_TYPE.PARAGRAPH)
        subheading_font = subheading_style.font
        subheading_font.name = 'Arial'
        subheading_font.size = Pt(12)
        subheading_font.bold = True
        subheading_style.paragraph_format.space_before = Pt(8)
        subheading_style.paragraph_format.space_after = Pt(4)
        
        # Body text style
        body_style = self.doc.styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
        body_font = body_style.font
        body_font.name = 'Arial'
        body_font.size = Pt(11)
        body_style.paragraph_format.space_after = Pt(6)
        body_style.paragraph_format.line_spacing = 1.15
    
    def add_header(self, title="Data Analysis Report", subtitle=None):
        """Add header section to the document."""
        title_para = self.doc.add_paragraph(title, style='CustomTitle')
        
        if subtitle:
            subtitle_para = self.doc.add_paragraph(subtitle, style='CustomHeading')
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        timestamp = datetime.now().strftime("%B %d, %Y at %I:%M %p")
        info_para = self.doc.add_paragraph(f"Generated on: {timestamp}", style='CustomBody')
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph("_" * 60, style='CustomBody').alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def parse_and_format_content(self, content):
        """Parse LLM response content and format it appropriately."""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Check for markdown table
            if '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
                table_lines = []
                j = i
                
                while j < len(lines) and '|' in lines[j].strip():
                    table_line = lines[j].strip()
                    if table_line:
                        table_lines.append(table_line)
                    j += 1
                
                if len(table_lines) >= 2:
                    if '-' in table_lines[1] or '=' in table_lines[1]:
                        self.add_table_from_markdown(table_lines)
                        i = j
                        continue
                    elif len(table_lines) >= 3 and ('-' in table_lines[2] or '=' in table_lines[2]):
                        self.add_table_from_markdown(table_lines)
                        i = j
                        continue
            
            # Check for headers
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                header_text = line.lstrip('#').strip()
                self.add_heading(header_text, level)
            
            # Check for bullet points
            elif line.startswith(('- ', '* ', '• ')):
                bullet_text = line[2:].strip()
                self.add_bullet_point(bullet_text)
            
            # Check for numbered lists
            elif line.split('.')[0].isdigit() and line.split('.', 1)[1].strip():
                numbered_text = line.split('.', 1)[1].strip()
                self.add_bullet_point(numbered_text, numbered=True)
            
            # Regular paragraph
            else:
                self.add_paragraph(line)
            
            i += 1
    
    def add_table_from_markdown(self, table_lines):
        """Convert markdown table to Word table."""
        if len(table_lines) < 2:
            return
        
        cleaned_lines = []
        for line in table_lines:
            cleaned = line.strip().strip('|')
            if cleaned and not all(c in '-=|: ' for c in cleaned):
                cleaned_lines.append(cleaned)
        
        if len(cleaned_lines) < 1:
            return
        
        header_row = [cell.strip() for cell in cleaned_lines[0].split('|')]
        data_rows = []
        
        for line in cleaned_lines[1:]:
            if line.strip():
                row = [cell.strip() for cell in line.split('|')]
                while len(row) < len(header_row):
                    row.append('')
                data_rows.append(row[:len(header_row)])
        
        if not data_rows:
            return
        
        try:
            table = self.doc.add_table(rows=1, cols=len(header_row))
            table.style = 'Table Grid'
            
            header_cells = table.rows[0].cells
            for i, header_text in enumerate(header_row):
                if i < len(header_cells):
                    header_cells[i].text = str(header_text)
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                    header_cells[i]._element.get_or_add_tcPr().append(
                        parse_xml(r'<w:shd {} w:fill="E7E6E6"/>'.format(nsdecls('w')))
                    )
            
            for row_data in data_rows:
                row_cells = table.add_row().cells
                for i, cell_text in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_text)
            
            self.doc.add_paragraph()
            
        except Exception:
            self.add_paragraph("Table Data:")
            self.add_paragraph(" | ".join(header_row))
            self.add_paragraph("-" * 50)
            for row in data_rows:
                self.add_paragraph(" | ".join(str(cell) for cell in row))
    
    def add_paragraph(self, text):
        """Add a paragraph to the document."""
        self.doc.add_paragraph(text, style='CustomBody')
    
    def add_heading(self, text, level=1):
        """Add a header to the document."""
        if level == 1:
            self.doc.add_paragraph(text, style='CustomHeading')
        else:
            self.doc.add_paragraph(text, style='CustomSubheading')
    
    def add_bullet_point(self, text, numbered=False):
        """Add a bullet point to the document."""
        if numbered:
            self.doc.add_paragraph(text, style='List Number')
        else:
            self.doc.add_paragraph(text, style='List Bullet')
    
    def save_document(self, filename=None, output_dir=None):
        """Save the document to a file."""
        if not filename:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"analysis_report_{timestamp}.docx"
        
        if output_dir:
            output_path = Path(output_dir) / filename
        else:
            output_path = Path(filename)
        
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        self.doc.save(str(output_path))
        return str(output_path)


class DataAnalysisTool:
    """Main application class that orchestrates the entire workflow."""
    
    def __init__(self):
        self.extractor = FileExtractor()
        self.structure_preview = StructurePreview()
        self.analyzer = None
        self.use_structure_template = False
    
    def display_banner(self):
        """Display welcome banner."""
        print("=" * 70)
        print("🚀 UNIFIED DATA ANALYSIS TOOL")
        print("   File Extraction + LLM Analysis + Professional Reports")
        print("=" * 70)
        print()
    
    def get_file_path(self):
        """Get file path from user input with validation."""
        while True:
            print("📁 FILE SELECTION")
            print("-" * 30)
            
            current_dir = Path.cwd()
            print(f"Current directory: {current_dir}")
            
            files = [f for f in current_dir.iterdir() if f.is_file()]
            if files:
                print("\nAvailable files in current directory:")
                for i, file in enumerate(files[:15], 1):
                    size = self.extractor.format_file_size(file.stat().st_size)
                    print(f"  {i}. {file.name} ({size})")
                if len(files) > 15:
                    print(f"  ... and {len(files) - 15} more files")
            
            print("\nOptions:")
            print("1. Enter file path manually")
            print("2. Use comprehensive_sales_data.csv (if available)")
            print("3. Use detailed_report_template.md (if available)")
            if files:
                print(f"4-{min(18, len(files) + 3)}. Select file by number from list above")
            print("99. Exit")
            
            choice = input("\nSelect option: ").strip()
            sys.stdout.flush()
            
            if choice == "1":
                file_path = input("Enter file path: ").strip().strip('"\'')
                sys.stdout.flush()
                if self.validate_file(file_path):
                    return file_path
                    
            elif choice == "2":
                sample_path = current_dir / "comprehensive_sales_data.csv"
                if sample_path.exists():
                    return str(sample_path)
                else:
                    print("❌ comprehensive_sales_data.csv not found in current directory")
                    
            elif choice == "3":
                template_path = current_dir / "detailed_report_template.md"
                if template_path.exists():
                    return str(template_path)
                else:
                    print("❌ detailed_report_template.md not found in current directory")
                    
            elif choice == "99":
                print("👋 Goodbye!")
                sys.exit(0)
                
            else:
                try:
                    file_num = int(choice)
                    if 4 <= file_num <= min(18, len(files) + 3):
                        file_index = file_num - 4
                        if file_index < len(files):
                            selected_file = files[file_index]
                            print(f"✅ Selected: {selected_file.name}")
                            return str(selected_file)
                        else:
                            print("❌ Invalid file number")
                    else:
                        print("❌ Invalid choice. Please try again.")
                except ValueError:
                    print("❌ Invalid choice. Please enter a number.")
    
    def setup_structure_preview(self):
        """Setup structure preview template from a selected file."""
        print("\n🏗️  STRUCTURE PREVIEW SETUP")
        print("-" * 40)
        print("Select a file to use as a structure template for all future analyses.")
        print("This will ensure consistent formatting across all reports.")
        
        while True:
            print("\nOptions:")
            print("1. Select a structure preview file")
            print("2. Skip structure preview (use default formatting)")
            print("3. View current structure template (if any)")
            
            choice = input("\nSelect option (1-3): ").strip()
            sys.stdout.flush()
            
            if choice == "1":
                preview_file = self.get_structure_preview_file()
                if preview_file:
                    try:
                        print(f"\n📋 Extracting structure from: {Path(preview_file).name}")
                        structure = self.structure_preview.extract_structure_from_file(preview_file)
                        
                        print("✅ Structure template extracted successfully!")
                        print(f"📊 Found {len(structure['sections'])} sections")
                        print(f"🎨 Organization style: {structure['organization_style']}")
                        
                        print("\n" + self.structure_preview.get_structure_summary())
                        
                        self.use_structure_template = True
                        return True
                        
                    except Exception as e:
                        print(f"❌ Error extracting structure: {str(e)}")
                        continue
                        
            elif choice == "2":
                print("📝 Using default formatting for analyses.")
                self.use_structure_template = False
                return True
                
            elif choice == "3":
                if self.structure_preview.structure_template:
                    print("\n" + self.structure_preview.get_structure_summary())
                else:
                    print("❌ No structure template currently loaded.")
                    
            else:
                print("❌ Invalid choice. Please try again.")
    
    def get_structure_preview_file(self):
        """Get structure preview file path from user input."""
        while True:
            print("\n📁 SELECT STRUCTURE PREVIEW FILE")
            print("-" * 35)
            
            current_dir = Path.cwd()
            print(f"Current directory: {current_dir}")
            
            files = [f for f in current_dir.iterdir() if f.is_file()]
            doc_files = [f for f in files if f.suffix.lower() in ['.txt', '.md', '.docx', '.doc']]
            
            if doc_files:
                print("\nAvailable document files:")
                for i, file in enumerate(doc_files, 1):
                    size = self.extractor.format_file_size(file.stat().st_size)
                    print(f"  {i}. {file.name} ({size})")
            
            print("\nOptions:")
            print("1. Enter file path manually")
            if doc_files:
                print(f"2-{len(doc_files) + 1}. Select file by number from list above")
            print("99. Back to main menu")
            
            choice = input("\nSelect option: ").strip()
            sys.stdout.flush()
            
            if choice == "1":
                file_path = input("Enter file path: ").strip().strip('"\'')
                sys.stdout.flush()
                if self.validate_file(file_path):
                    return file_path
                    
            elif choice == "99":
                return None
                
            else:
                try:
                    file_num = int(choice)
                    if 2 <= file_num <= len(doc_files) + 1:
                        file_index = file_num - 2
                        if file_index < len(doc_files):
                            selected_file = doc_files[file_index]
                            print(f"✅ Selected: {selected_file.name}")
                            return str(selected_file)
                        else:
                            print("❌ Invalid file number")
                    else:
                        print("❌ Invalid choice. Please try again.")
                except ValueError:
                    print("❌ Invalid choice. Please enter a number.")
            
            print()
    
    def validate_file(self, file_path):
        """Validate if file exists and is supported."""
        path = Path(file_path)
        
        if not path.exists():
            print(f"❌ File not found: {file_path}")
            return False
            
        if not path.is_file():
            print(f"❌ Path is not a file: {file_path}")
            return False
            
        size = path.stat().st_size
        if size > 10 * 1024 * 1024:  # 10MB
            print(f"⚠️  Large file detected ({self.extractor.format_file_size(size)})")
            confirm = input("Continue anyway? (y/n): ").lower()
            if confirm != 'y':
                return False
        
        print(f"✅ File validated: {path.name} ({self.extractor.format_file_size(size)})")
        return True
    
    def get_api_settings(self):
        """Get API provider and credentials from user."""
        print("\n🤖 LLM API CONFIGURATION")
        print("-" * 30)
        
        print("Available API providers:")
        print("1. OpenAI (GPT-4)")
        print("2. Anthropic (Claude)")
        print("3. Ollama (Local)")
        
        while True:
            choice = input("\nSelect API provider (1-3): ").strip()
            
            if choice == "1":
                provider = "openai"
                api_key = self.get_api_key("OpenAI")
                break
            elif choice == "2":
                provider = "anthropic"
                api_key = self.get_api_key("Anthropic")
                break
            elif choice == "3":
                provider = "ollama"
                api_key = None
                print("✅ Using local Ollama (no API key required)")
                break
            else:
                print("❌ Invalid choice. Please try again.")
        
        return provider, api_key
    
    def get_api_key(self, provider_name):
        """Get API key from user or environment."""
        env_key = os.getenv('API_KEY')
        if env_key:
            print(f"✅ Using API key from environment variable")
            return env_key
        
        print(f"\n🔑 {provider_name} API Key Required")
        print("You can:")
        print("1. Enter API key now")
        print("2. Set API_KEY environment variable")
        print("3. Skip (will cause errors)")
        
        choice = input("Choose option (1-3): ").strip()
        
        if choice == "1":
            api_key = input(f"Enter {provider_name} API key: ").strip()
            if api_key:
                return api_key
            else:
                print("❌ Empty API key provided")
                return None
        elif choice == "2":
            print(f"💡 Set environment variable: export API_KEY='your-{provider_name.lower()}-key'")
            return None
        else:
            print("⚠️  Proceeding without API key (may cause errors)")
            return None
    
    def extract_and_preview(self, file_path):
        """Extract file content and show preview."""
        print(f"\n📄 EXTRACTING FILE CONTENT")
        print("-" * 30)
        
        try:
            file_info = self.extractor.get_file_info(file_path)
            
            if "error" in file_info:
                print(f"❌ {file_info['error']}")
                return None
            
            print(f"File: {file_info['filename']}")
            print(f"Size: {file_info['size_readable']}")
            print(f"Type: {file_info['extension']}")
            
            content = self.extractor.extract_file_content(file_path)
            
            if content.startswith("Error"):
                print(f"❌ {content}")
                return None
            
            print(f"✅ Successfully extracted {len(content)} characters")
            
            preview_length = 300
            if len(content) > preview_length:
                preview = content[:preview_length] + "..."
                print(f"\n📋 Content Preview (first {preview_length} chars):")
            else:
                preview = content
                print(f"\n📋 Full Content:")
            
            print("-" * 50)
            print(preview)
            print("-" * 50)
            
            return content
            
        except Exception as e:
            print(f"❌ Error extracting file content: {str(e)}")
            return None
    
    def analyze_with_llm(self, content, file_path, provider, api_key):
        """Send content to LLM for analysis with optional structure template."""
        print(f"\n🧠 LLM ANALYSIS")
        print("-" * 30)
        
        try:
            structure_prompt = ""
            if self.use_structure_template and self.structure_preview.structure_template:
                structure_prompt = self.structure_preview.generate_structure_prompt()
                print("📋 Using structure template for consistent formatting")
            
            self.analyzer = LLMAnalyzer(api_key=api_key, api_provider=provider)
            
            print(f"🔗 Using {provider.upper()} API")
            print("🚀 Sending content for analysis...")
            
            analysis = self.analyzer.analyze_content(
                content, 
                file_path,
                structure_template=structure_prompt
            )
            
            if analysis:
                print("\n📊 ANALYSIS COMPLETE")
                print("=" * 60)
                print(analysis)
                print("=" * 60)
                return analysis
            else:
                print("❌ No analysis received from LLM")
                return None
                
        except Exception as e:
            print(f"❌ LLM Analysis Error: {str(e)}")
            return None
    
    def generate_docx_report(self, analysis, file_path):
        """Generate a DOCX report from the LLM analysis."""
        print(f"\n📄 DOCX REPORT GENERATION")
        print("-" * 30)
        
        if not DOCX_AVAILABLE:
            print("⚠️  DOCX generation not available. Install python-docx: pip install python-docx")
            return None
        
        try:
            print("Would you like to generate a formatted DOCX report?")
            print("1. Yes, generate DOCX report")
            print("2. No, skip DOCX generation")
            
            choice = input("\nSelect option (1-2): ").strip()
            
            if choice == "1":
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                original_name = Path(file_path).stem
                docx_filename = f"analysis_report_{original_name}_{timestamp}.docx"
                
                generator = DocxGenerator()
                
                file_name = Path(file_path).name
                generator.add_header("Comprehensive Data Analysis Report", f"Analysis of {file_name}")
                generator.parse_and_format_content(analysis)
                
                docx_path = generator.save_document(docx_filename)
                
                print(f"✅ DOCX report generated: {docx_path}")
                return docx_path
                
            else:
                print("⏭️  Skipping DOCX generation")
                return None
                
        except Exception as e:
            print(f"❌ Error generating DOCX: {str(e)}")
            return None
    
    def save_results(self, analysis, file_path, content, docx_path=None):
        """Save analysis results and original content."""
        print(f"\n💾 SAVING RESULTS")
        print("-" * 30)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        original_name = Path(file_path).stem
        
        report_filename = f"analysis_report_{original_name}_{timestamp}.txt"
        
        try:
            with open(report_filename, 'w', encoding='utf-8') as f:
                f.write("FILE ANALYSIS REPORT\n")
                f.write("=" * 60 + "\n")
                f.write(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Original File: {file_path}\n")
                f.write(f"Content Length: {len(content)} characters\n")
                if docx_path:
                    f.write(f"DOCX Report: {docx_path}\n")
                f.write("=" * 60 + "\n\n")
                f.write("ANALYSIS RESULTS:\n")
                f.write("-" * 30 + "\n")
                f.write(analysis)
                f.write("\n\n" + "=" * 60 + "\n")
                f.write("ORIGINAL CONTENT:\n")
                f.write("-" * 30 + "\n")
                f.write(content)
            
            print(f"✅ Report saved: {report_filename}")
            
            summary_filename = f"summary_{original_name}_{timestamp}.txt"
            with open(summary_filename, 'w', encoding='utf-8') as f:
                f.write(f"Analysis Summary - {original_name}\n")
                f.write("=" * 40 + "\n")
                f.write(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"File: {Path(file_path).name}\n")
                f.write(f"Size: {len(content)} characters\n")
                if docx_path:
                    f.write(f"DOCX: {Path(docx_path).name}\n")
                f.write("\nAnalysis:\n")
                f.write(analysis[:500] + "..." if len(analysis) > 500 else analysis)
            
            print(f"✅ Summary saved: {summary_filename}")
            if docx_path:
                print(f"✅ DOCX report: {docx_path}")
            
        except Exception as e:
            print(f"⚠️  Could not save files: {str(e)}")
    
    def run(self):
        """Main workflow execution."""
        self.display_banner()
        
        try:
            # Step 0: Setup structure preview (optional)
            if not self.setup_structure_preview():
                return
            
            # Step 1: File selection
            file_path = self.get_file_path()
            
            # Step 2: Extract and preview content
            content = self.extract_and_preview(file_path)
            if not content:
                print("❌ Cannot proceed without file content")
                return
            
            # Step 3: API configuration
            provider, api_key = self.get_api_settings()
            
            # Step 4: LLM analysis
            analysis = self.analyze_with_llm(content, file_path, provider, api_key)
            if not analysis:
                print("❌ Cannot proceed without analysis")
                return
            
            # Step 5: Generate DOCX report (optional)
            docx_path = self.generate_docx_report(analysis, file_path)
            
            # Step 6: Save results
            self.save_results(analysis, file_path, content, docx_path)
            
            print("\n🎉 WORKFLOW COMPLETED SUCCESSFULLY!")
            print("=" * 50)
            print("✅ File analyzed")
            print("✅ Report generated")
            print("✅ Results saved")
            if docx_path:
                print("✅ DOCX report created")
            
        except KeyboardInterrupt:
            print("\n\n⚠️  Workflow interrupted by user")
        except Exception as e:
            print(f"\n❌ Unexpected error: {str(e)}")
            print("Please check your inputs and try again.")


def main():
    """Main entry point."""
    print("🔧 Checking dependencies...")
    
    missing_deps = []
    if not DOCX_AVAILABLE:
        missing_deps.append("python-docx (for DOCX generation)")
    if not PANDAS_AVAILABLE:
        missing_deps.append("pandas (for Excel file support)")
    if not PDF_AVAILABLE:
        missing_deps.append("PyPDF2 (for PDF file support)")
    
    if missing_deps:
        print("⚠️  Optional dependencies missing:")
        for dep in missing_deps:
            print(f"   - {dep}")
        print("\nInstall with: pip install python-docx pandas PyPDF2")
        print("The tool will work with reduced functionality.\n")
    
    # Initialize and run the tool
    tool = DataAnalysisTool()
    tool.run()


if __name__ == "__main__":
    main()