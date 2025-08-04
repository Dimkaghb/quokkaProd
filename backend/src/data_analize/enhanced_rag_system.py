"""
Enhanced RAG System with persistent Chroma vector store and optimized processing.
Addresses memory issues with large document collections.
"""

import os
import logging
import uuid
import pandas as pd
from typing import List, Dict, Any, Optional
from pathlib import Path

from dotenv import load_dotenv
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_core.messages import HumanMessage, AIMessage
from langchain.memory import ConversationBufferWindowMemory
from langchain.schema import BaseMessage

load_dotenv()

logger = logging.getLogger(__name__)


class EnhancedRAGSystem:
    """
    Enhanced RAG system with persistent vector storage and optimized memory usage.
    """
    
    def __init__(
        self,
        collection_name: str = "default",
        persist_directory: str = None,
        model_name: str = "gpt-4o-mini",
        chunk_size: int = 2000,  # Increased for better context
        chunk_overlap: int = 400  # Increased overlap for better continuity
    ):
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        if not self.openai_api_key:
            raise ValueError("OPENAI_API_KEY not found in environment variables")
        
        self.model_name = model_name
        self.collection_name = collection_name
        self.persist_directory = persist_directory or self._get_default_persist_dir()
        
        # Ensure persist directory exists
        Path(self.persist_directory).mkdir(parents=True, exist_ok=True)
        
        self.embeddings = OpenAIEmbeddings(
            openai_api_key=self.openai_api_key,
            model="text-embedding-3-small"  # More efficient embedding model
        )
        
        self.llm = ChatOpenAI(
            model=self.model_name,
            temperature=0.7,
            openai_api_key=self.openai_api_key
        )
        
        # Use persistent Chroma vector store
        self.vector_store = Chroma(
            collection_name=self.collection_name,
            embedding_function=self.embeddings,
            persist_directory=self.persist_directory
        )
        
        # Optimized text splitter for better chunks
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", "? ", "! ", " ", ""]
        )
        
        self.memory = ConversationBufferWindowMemory(
            k=15,  # Increased for better conversation context
            return_messages=True,
            memory_key="chat_history"
        )
        
        self.rag_chain = None
        self._setup_rag_chain()
        
        logger.info(f"Enhanced RAG system initialized with collection: {self.collection_name}")
    
    def _get_default_persist_dir(self) -> str:
        """Get default persistence directory."""
        current_dir = Path(__file__).parent
        return str(current_dir.parent.parent / "data" / "chroma_db")
    
    def _setup_rag_chain(self):
        """Setup the RAG chain with enhanced prompting."""
        retriever = self.vector_store.as_retriever(
            search_type="similarity",
            search_kwargs={"k": 5}  # Increased for better context
        )
        
        prompt_template = """You are an expert data analyst and document assistant with deep expertise in analysis, pattern detection, and strategic insights.

Based on the provided context from documents, provide detailed, comprehensive, and professional responses that directly address the user's query.

Your responses should be:
- Direct and focused on answering the specific question
- Detailed with thorough analysis and insights
- Professional in tone and presentation
- Evidence-based using the provided context
- Practical with actionable information when relevant
- Well-structured with clear sections when appropriate

Context from documents:
{context}

Conversation History:
{chat_history}

Current Question: {question}

If the context doesn't contain sufficient information to answer the question, clearly state what information is missing and what you can provide based on the available context.

Provide a detailed, professional response:"""
        
        prompt = ChatPromptTemplate.from_template(prompt_template)
        
        def format_docs(docs):
            """Enhanced document formatting with metadata."""
            formatted_docs = []
            for doc in docs:
                content = doc.page_content.strip()
                source = doc.metadata.get("source", "Unknown")
                formatted_docs.append(f"[Source: {source}]\n{content}")
            return "\n\n---\n\n".join(formatted_docs)
        
        def format_chat_history(messages):
            """Enhanced chat history formatting."""
            if not messages:
                return "No previous conversation."
            
            formatted = []
            for msg in messages[-10:]:  # Last 10 messages for context
                if isinstance(msg, HumanMessage):
                    formatted.append(f"Human: {msg.content}")
                elif isinstance(msg, AIMessage):
                    formatted.append(f"Assistant: {msg.content}")
            return "\n".join(formatted)
        
        self.rag_chain = (
            {
                "context": retriever | format_docs,
                "question": RunnablePassthrough(),
                "chat_history": lambda x: format_chat_history(
                    self.memory.chat_memory.messages
                )
            }
            | prompt
            | self.llm
            | StrOutputParser()
        )
    
    async def add_document_async(
        self,
        file_path: str,
        metadata: Optional[Dict[str, Any]] = None,
        chunk_documents: bool = True
    ) -> Dict[str, Any]:
        """
        Add a document to the vector store asynchronously.
        
        Args:
            file_path: Path to the document file
            metadata: Optional metadata for the document
            chunk_documents: Whether to chunk the document
            
        Returns:
            Result information
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            # Read file content based on type
            content = await self._read_file_content_async(file_path)
            
            if not content.strip():
                logger.warning(f"Empty content in file: {file_path}")
                return {"success": False, "error": "Empty file content"}
            
            # Prepare metadata
            if metadata is None:
                metadata = {}
            
            from datetime import datetime
            metadata.update({
                "source": file_path,
                "file_name": Path(file_path).name,
                "doc_id": str(uuid.uuid4()),
                "processing_timestamp": datetime.now().isoformat()
            })
            
            # Create document
            document = Document(page_content=content, metadata=metadata)
            
            # Split into chunks if requested
            if chunk_documents:
                chunks = self.text_splitter.split_documents([document])
                logger.info(f"Split document into {len(chunks)} chunks")
            else:
                chunks = [document]
            
            # Add to vector store
            doc_ids = self.vector_store.add_documents(chunks)
            
            # Persist the changes
            self.vector_store.persist()
            
            logger.info(f"Added document {file_path} with {len(chunks)} chunks")
            
            return {
                "success": True,
                "chunks_added": len(chunks),
                "doc_ids": doc_ids,
                "file_path": file_path
            }
            
        except Exception as e:
            logger.error(f"Error adding document {file_path}: {e}")
            return {"success": False, "error": str(e)}
    
    async def _read_file_content_async(self, file_path: str) -> str:
        """Read file content asynchronously based on file type."""
        import aiofiles
        
        file_path_obj = Path(file_path)
        file_ext = file_path_obj.suffix.lower()
        
        try:
            if file_ext in ['.txt', '.md']:
                async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                    return await f.read()
            
            elif file_ext == '.pdf':
                # Use PyMuPDF for PDF reading
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            
            elif file_ext in ['.docx', '.doc']:
                # Use python-docx for Word documents
                from docx import Document as DocxDocument
                doc = DocxDocument(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            
            elif file_ext in ['.csv', '.xlsx', '.xls']:
                # For data files, create a summary
                import pandas as pd
                df = pd.read_csv(file_path) if file_ext == '.csv' else pd.read_excel(file_path)
                
                summary = f"""Data File Summary for {file_path_obj.name}:

Shape: {df.shape[0]} rows, {df.shape[1]} columns

Columns: {', '.join(df.columns.tolist())}

Data Types:
{df.dtypes.to_string()}

Basic Statistics:
{df.describe(include='all').to_string()}

First 5 rows:
{df.head().to_string()}
"""
                return summary
            
            else:
                # Fallback: try to read as text
                async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return await f.read()
                    
        except Exception as e:
            logger.error(f"Error reading file {file_path}: {e}")
            raise
    
    def query(self, question: str) -> Dict[str, Any]:
        """
        Query the RAG system.
        
        Args:
            question: User question
            
        Returns:
            Response with answer, sources, and metadata
        """
        try:
            # Get response from RAG chain
            response = self.rag_chain.invoke(question)
            
            # Add to conversation memory
            self.memory.chat_memory.add_user_message(question)
            self.memory.chat_memory.add_ai_message(response)
            
            # Get source documents for citation
            retrieved_docs = self.vector_store.similarity_search(question, k=5)
            sources = []
            for doc in retrieved_docs:
                source_info = {
                    "source": doc.metadata.get("source", "Unknown"),
                    "file_name": doc.metadata.get("file_name", "Unknown"),
                    "relevance_score": getattr(doc, 'relevance_score', None)
                }
                if source_info not in sources:
                    sources.append(source_info)
            
            return {
                "answer": response,
                "sources": sources,
                "conversation_id": len(self.memory.chat_memory.messages) // 2,
                "document_count": len(retrieved_docs)
            }
        
        except Exception as e:
            logger.error(f"Error processing query: {e}")
            return {
                "answer": "I apologize, but I encountered an error processing your question. Please try again.",
                "sources": [],
                "conversation_id": None,
                "error": str(e)
            }
    
    def get_collection_stats(self) -> Dict[str, Any]:
        """Get statistics about the vector store collection."""
        try:
            # Get collection info
            collection = self.vector_store._collection
            count = collection.count()
            
            return {
                "collection_name": self.collection_name,
                "document_count": count,
                "persist_directory": self.persist_directory
            }
        except Exception as e:
            logger.error(f"Error getting collection stats: {e}")
            return {"error": str(e)}
    
    def clear_collection(self) -> bool:
        """Clear all documents from the collection."""
        try:
            # Get all document IDs and delete them
            collection = self.vector_store._collection
            result = collection.get()
            if result and result.get('ids'):
                collection.delete(ids=result['ids'])
                self.vector_store.persist()
                logger.info(f"Cleared {len(result['ids'])} documents from collection")
            return True
        except Exception as e:
            logger.error(f"Error clearing collection: {e}")
            return False
    
    def clear_memory(self):
        """Clear conversation memory."""
        self.memory.clear()
        logger.info("Conversation memory cleared")
    
    def get_conversation_history(self) -> List[Dict[str, str]]:
        """Get conversation history."""
        messages = self.memory.chat_memory.messages
        history = []
        
        for i in range(0, len(messages), 2):
            if i + 1 < len(messages):
                history.append({
                    "question": messages[i].content,
                    "answer": messages[i + 1].content,
                    "timestamp": getattr(messages[i], 'timestamp', None)
                })
        
        return history


# Global instance for the application
_rag_instance = None

def get_rag_system(collection_name: str = "default") -> EnhancedRAGSystem:
    """Get a global RAG system instance."""
    global _rag_instance
    if _rag_instance is None or _rag_instance.collection_name != collection_name:
        _rag_instance = EnhancedRAGSystem(collection_name=collection_name)
    return _rag_instance